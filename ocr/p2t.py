from enum import Enum
from io import BytesIO
from pathlib import Path
from typing import Literal

import latex2mathml.converter
import latex2mathml.exceptions
from docx import Document
from lxml import etree
from PIL.Image import Image as ImageType
from pix2text import Pix2Text

from ocr.sanitiser import Sanitiser

FILE_DIR = Path(__file__).resolve().parent
BASE_DIR = FILE_DIR.parent


class FormatConverter:
	def __init__(self):
		self.mathml_to_omml = etree.XSLT(etree.parse(BASE_DIR / 'data' / 'MML2OMML.XSL'))  # NOSONAR

	def latex_to_mathml(self, latex: str) -> str:
		mathml = latex2mathml.converter.convert(latex)
		mathml = mathml.replace(' display="inline"', '')
		return mathml

	def convert_latex_to_mathml(self, latex: str) -> str:
		return self.latex_to_mathml(latex)

	def convert_mathml_to_xml(self, mathml: str) -> etree._ElementTree:
		return etree.fromstring(mathml)

	def convert_mathmlxml_to_omml(self, mathml: etree._ElementTree) -> etree._ElementTree:
		return self.mathml_to_omml(mathml)

	def convert_omml_to_docx_element(self, omml: etree._ElementTree) -> etree._Element:
		return omml.getroot()

	def try_convert_latex_to_mathml(self, latex: str) -> str:
		try:
			return formatter.convert_latex_to_mathml(latex)
		except Exception:
			return latex

	def try_convert_latex_to_omml(self, latex: str) -> str:
		try:
			output = formatter.convert_latex_to_mathml(latex)
			output = formatter.convert_mathml_to_xml(output)
			output = formatter.convert_mathmlxml_to_omml(output)
			return str(output)
		except Exception:
			return latex


class P2TAnalyser:
	model = Pix2Text

	def __init__(self, languages=('en',)):
		self.model = Pix2Text.from_config(languages=languages)

	def analyse(
		self,
		image: ImageType,
		type: Literal['text', 'formula', 'text_formula', 'page', 'pdf'] = 'text_formula',
	):
		result = self.model.recognize(image, file_type=type)
		return result


class P2TInput(Enum):
	TEXT = 'text'
	FORMULA = 'formula'
	TEXT_FORMULA = 'text_formula'
	PAGE = 'page'
	PDF = 'pdf'


class P2TOutput(Enum):
	LATEX = 'latex'
	MATHML = 'mathml'
	OMML = 'omml'
	DOCX = 'docx'


analyser = P2TAnalyser()
formatter = FormatConverter()
sanitiser = Sanitiser()


def convert_output(results: list[str], output_type: P2TOutput) -> list[str] | BytesIO:
	match output_type:
		case P2TOutput.LATEX:
			return results

		case P2TOutput.MATHML:
			return [formatter.try_convert_latex_to_mathml(result) for result in results]

		case P2TOutput.OMML:
			return [
				formatter.try_convert_latex_to_omml(result)
				for result in results
				if result.startswith(Sanitiser.MATH_BEGIN)
			]

		case P2TOutput.DOCX:
			document = Document()
			p = document.add_paragraph()

			for result in results:
				if result.startswith(Sanitiser.MATH_BEGIN):
					result = formatter.convert_latex_to_mathml(result)
					result = formatter.convert_mathml_to_xml(result)
					result = formatter.convert_mathmlxml_to_omml(result)
					result = formatter.convert_omml_to_docx_element(result)
					p._element.append(result)
				else:
					p.add_run(result)

			file_stream = BytesIO()
			document.save(file_stream)

			return file_stream


def analyse_p2t(
	image: ImageType | Path,
	input_type: P2TInput = P2TInput.TEXT_FORMULA,
	output_type: P2TOutput | list[P2TOutput] = P2TOutput.LATEX,
) -> list[str] | BytesIO | dict[str, list[str] | BytesIO]:
	results = analyser.analyse(image, input_type.value)

	match input_type.value:
		case P2TInput.TEXT_FORMULA.value | P2TInput.FORMULA.value:
			results = sanitiser.clean_mix_output(results)
		case P2TInput.TEXT.value:
			results = [results]
		case P2TInput.PDF.value:
			results.to_markdown('output-pdf-md')
		case _:
			raise NotImplementedError('Input Not Implemented!')

	if isinstance(output_type, list):
		outputs = {out_type.value: convert_output(results, out_type) for out_type in output_type}
		return outputs

	return convert_output(results, output_type)
